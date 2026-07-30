"""Tests for parsers, chunking, and indexing pipeline (M2)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

from docuwizard.db import connect
from docuwizard.ingest import chunking, parsers, store
from docuwizard.ingest.pipeline import index_file, index_project_files
from docuwizard.ingest.segments import TextSegment
from docuwizard.models import FileStatus
from docuwizard.services import files as file_service
from docuwizard.services import projects as project_service


def test_parse_text_line_metadata(tmp_path: Path) -> None:
    path = tmp_path / "note.txt"
    path.write_text("첫번째\n\n두번째\n", encoding="utf-8")
    segments = parsers.parse_text(path)
    assert [(s.text, s.line_start, s.line_end) for s in segments] == [
        ("첫번째", 1, 1),
        ("두번째", 3, 3),
    ]


def test_parse_docx(tmp_path: Path) -> None:
    path = tmp_path / "doc.docx"
    document = Document()
    document.add_paragraph("제출 서류")
    document.add_paragraph("마감일 안내")
    document.save(path)
    segments = parsers.parse_docx(path)
    assert [s.text for s in segments] == ["제출 서류", "마감일 안내"]
    assert segments[0].line_start == 1


def test_parse_xlsx(tmp_path: Path) -> None:
    path = tmp_path / "sheet.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "일정"
    ws["A1"] = "항목"
    ws["B1"] = "마감"
    ws["A2"] = "제안서"
    ws["B2"] = "2026-08-01"
    wb.save(path)
    segments = parsers.parse_xlsx(path)
    assert any(s.sheet == "일정" and "제안서" in s.text for s in segments)
    assert any(s.cell_range for s in segments)


def _write_hwpx(path: Path, sections: list[list[str]]) -> None:
    """Create a minimal HWPX archive with the given paragraphs per section."""
    import zipfile

    ns = 'xmlns:hp="http://www.hancom.co.kr/hwpml/2011/paragraph"'
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("mimetype", "application/hwp+zip")
        for idx, paragraphs in enumerate(sections):
            body = "".join(
                f"<hp:p><hp:run><hp:t>{text}</hp:t></hp:run></hp:p>"
                for text in paragraphs
            )
            zf.writestr(
                f"Contents/section{idx}.xml",
                f'<?xml version="1.0" encoding="UTF-8"?><hp:sec {ns}>{body}</hp:sec>',
            )


def test_parse_hwpx(tmp_path: Path) -> None:
    path = tmp_path / "doc.hwpx"
    _write_hwpx(path, [["제출 마감일 안내", "", "평가 기준"]])
    segments = parsers.parse_hwpx(path)
    assert [(s.text, s.line_start) for s in segments] == [
        ("제출 마감일 안내", 1),
        ("평가 기준", 2),
    ]
    # Single section → no page metadata.
    assert segments[0].page is None


def test_parse_hwpx_multi_section_pages(tmp_path: Path) -> None:
    path = tmp_path / "multi.hwpx"
    _write_hwpx(path, [["1장 내용"], ["2장 내용"]])
    segments = parsers.parse_hwpx(path)
    assert [(s.text, s.page) for s in segments] == [
        ("1장 내용", 1),
        ("2장 내용", 2),
    ]


def test_parse_hwpx_rejects_non_zip(tmp_path: Path) -> None:
    import pytest

    path = tmp_path / "broken.hwpx"
    path.write_bytes(b"not a zip")
    with pytest.raises(parsers.ParseError):
        parsers.parse_hwpx(path)


def test_index_hwpx_file(tmp_path: Path) -> None:
    from fakes import FakeOllama

    project = project_service.create_project("한글")
    src = tmp_path / "notice.hwpx"
    _write_hwpx(src, [["제안서 마감일은 8월 1일입니다."]])
    added = file_service.add_files(project.id, [src])[0]
    count = index_file(project.id, added, embedder=FakeOllama())
    assert count >= 1
    rows = store.list_chunks_for_file(added.id)
    assert "마감일" in rows[0]["text"]


def test_parse_image_ocr(tmp_path: Path, monkeypatch) -> None:
    import pytesseract
    from PIL import Image

    path = tmp_path / "scan.png"
    Image.new("RGB", (60, 20), "white").save(path)
    monkeypatch.setattr(
        pytesseract,
        "image_to_string",
        lambda img, lang=None: "마감일 안내\n\n8월 1일\n",
    )
    segments = parsers.parse_image(path)
    assert [(s.text, s.line_start) for s in segments] == [
        ("마감일 안내", 1),
        ("8월 1일", 3),
    ]


def test_parse_image_missing_tesseract(tmp_path: Path, monkeypatch) -> None:
    import pytesseract
    import pytest
    from PIL import Image

    path = tmp_path / "scan.png"
    Image.new("RGB", (60, 20), "white").save(path)

    def raise_missing(img, lang=None):
        raise pytesseract.TesseractNotFoundError()

    monkeypatch.setattr(pytesseract, "image_to_string", raise_missing)
    with pytest.raises(parsers.ParseError, match="Tesseract"):
        parsers.parse_image(path)


def test_add_files_skips_duplicate_content(tmp_path: Path) -> None:
    project = project_service.create_project("중복")
    a = tmp_path / "a.txt"
    a.write_text("같은 내용", encoding="utf-8")
    b = tmp_path / "b.txt"
    b.write_text("같은 내용", encoding="utf-8")
    c = tmp_path / "c.txt"
    c.write_text("다른 내용", encoding="utf-8")

    added = file_service.add_files(project.id, [a, b, c])
    assert [f.original_name for f in added] == ["a.txt", "c.txt"]
    assert all(f.content_hash for f in added)

    # Re-adding an identical file is skipped too.
    assert file_service.add_files(project.id, [a]) == []
    assert len(file_service.list_files(project.id)) == 2


def test_index_embeds_in_batches(tmp_path: Path, monkeypatch) -> None:
    import copy

    from docuwizard.config import DEFAULT_SETTINGS
    from fakes import FakeOllama

    calls: list[int] = []

    class CountingOllama(FakeOllama):
        def embed(self, texts: list[str]) -> list[list[float]]:
            calls.append(len(texts))
            return super().embed(texts)

    settings = copy.deepcopy(DEFAULT_SETTINGS)
    settings["rag"]["embed_batch_size"] = 2
    monkeypatch.setattr("docuwizard.ingest.pipeline.load_settings", lambda: settings)

    project = project_service.create_project("배치")
    src = tmp_path / "long.txt"
    src.write_text("\n".join(f"{i}번째 줄 " + "가나다라" * 40 for i in range(30)), encoding="utf-8")
    added = file_service.add_files(project.id, [src])[0]
    count = index_file(project.id, added, embedder=CountingOllama())

    assert count > 2
    assert len(calls) >= 2
    assert all(size <= 2 for size in calls)
    assert sum(calls) == count


def test_count_stale_embeddings(tmp_path: Path) -> None:
    from docuwizard.rag import vectors
    from fakes import FakeOllama

    project = project_service.create_project("모델변경")
    src = tmp_path / "a.txt"
    src.write_text("마감일은 금요일입니다.", encoding="utf-8")
    added = file_service.add_files(project.id, [src])[0]
    index_file(project.id, added, embedder=FakeOllama())

    assert vectors.count_stale_embeddings(project.id, "fake-embed") == 0
    assert vectors.count_stale_embeddings(project.id, "새-모델") >= 1


def test_parse_pdf(tmp_path: Path) -> None:
    path = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with path.open("wb") as f:
        writer.write(f)
    # Blank PDF yields no text segments; ensure it does not crash.
    assert parsers.parse_pdf(path) == []


def test_chunk_segments_respects_size() -> None:
    segments = [TextSegment(text="가" * 50, line_start=i, line_end=i) for i in range(1, 21)]
    chunks = chunking.chunk_segments(segments, chunk_size=120, chunk_overlap=20)
    assert chunks
    assert all(len(c.text) <= 120 for c in chunks)
    assert chunks[0].line_start is not None


def test_index_text_file_writes_chunks(tmp_path: Path) -> None:
    from fakes import FakeOllama

    project = project_service.create_project("인덱싱")
    src = tmp_path / "guide.md"
    src.write_text("# 안내\n마감일은 다음 주 금요일입니다.\n", encoding="utf-8")
    added = file_service.add_files(project.id, [src])[0]

    count = index_file(project.id, added, embedder=FakeOllama())
    assert count >= 1
    refreshed = file_service.list_files(project.id)[0]
    assert refreshed.status == FileStatus.READY
    assert store.count_chunks(project.id) == count
    rows = store.list_chunks_for_file(added.id)
    assert rows[0]["text"]


def test_index_failed_retry_and_delete_clears_chunks(tmp_path: Path) -> None:
    from fakes import FakeOllama

    project = project_service.create_project("재시도")
    good = tmp_path / "ok.txt"
    good.write_text("유효한 내용", encoding="utf-8")
    unsupported = tmp_path / "image.webp"
    unsupported.write_bytes(b"webp")
    files = file_service.add_files(project.id, [good, unsupported])
    ok, failed = index_project_files(project.id, embedder=FakeOllama())
    assert ok == 1
    assert failed == 1

    statuses = {f.original_name: f.status for f in file_service.list_files(project.id)}
    assert statuses["ok.txt"] == FileStatus.READY
    assert statuses["image.webp"] == FileStatus.FAILED

    file_id = next(f.id for f in files if f.original_name == "ok.txt")
    assert store.list_chunks_for_file(file_id)
    file_service.delete_file(project.id, file_id)
    assert store.list_chunks_for_file(file_id) == []


def test_schema_tables_exist() -> None:
    with connect() as conn:
        tables = {
            row[0]
            for row in conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table'"
            ).fetchall()
        }
    assert {"projects", "files", "chunks", "conversations", "messages", "embeddings"} <= tables
